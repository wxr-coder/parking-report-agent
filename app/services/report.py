import re
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING

from app.services.agent import NarrativeSections
from app.services.metrics import MetricsResult, money_str

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from docx.text.paragraph import Paragraph


def render_report(
    *,
    template_path: Path,
    output_path: Path,
    metrics: MetricsResult,
    narrative: NarrativeSections,
    chart_path: Path,
) -> Path:
    from docx import Document

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document(template_path)
    _fill_key_metrics_table(document, metrics)
    _replace_paragraphs(document, metrics, narrative)
    _replace_section_body(document, "四、补充观察", _bullet_text(narrative.observations))
    _replace_section_body(document, "五、结论与建议", _bullet_text(narrative.recommendations))
    _replace_chart_section(document, chart_path)
    document.save(output_path)
    return output_path


def _fill_key_metrics_table(document: "DocxDocument", metrics: MetricsResult) -> None:
    values = {
        "1. 总交易笔数": f"{metrics.total_transactions} 笔",
        "2. 应收总金额（元）": f"¥ {money_str(metrics.expected_amount)}",
        "3. 实收总金额（元）": f"¥ {money_str(metrics.paid_amount)}",
        "4. 实际抵扣总额（元）": f"¥ {money_str(metrics.actual_discount_amount)}",
        "5. 实收率（%）": f"{metrics.collection_rate_percent:.1f}%",
        "6. 主要支付方式": f"{metrics.main_payment_method} ×{metrics.main_payment_method_count}",
    }
    for table in document.tables:
        for row in table.rows:
            if not row.cells:
                continue
            key = row.cells[0].text.strip()
            if key in values and len(row.cells) > 1:
                row.cells[1].text = values[key]


def _replace_paragraphs(
    document: "DocxDocument",
    metrics: MetricsResult,
    narrative: NarrativeSections,
) -> None:
    report_period = _report_period(metrics)
    replacements = {
        "数据周期：【起始日期 – 结束日期】": f"数据周期：{report_period}",
        "生成时间：【时间戳】": f"生成时间：{datetime.now():%Y-%m-%d %H:%M:%S}",
        "【叙述：各支付方式与支付渠道（线上支付 / 出口贴码）的分布情况，指出占主导的方式与渠道。】": narrative.payment_and_channel,
        "【根据“收费时间 − 进车时间”推算停车时长。报告平均值 / 分布情况及明显的长时停车异常项。可选择再加入一张图表（按小时或时长区间）。】": narrative.parking_duration,
    }
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        replacement = replacements.get(text)
        if replacement is not None:
            _set_paragraph_text(paragraph, replacement)


def _replace_chart_section(document: "DocxDocument", chart_path: Path) -> None:
    from docx.shared import Inches

    paragraphs = document.paragraphs
    title_index = _find_paragraph_index(paragraphs, "支付方式分布图")
    if title_index is None:
        document.add_picture(str(chart_path), width=Inches(5.8))
        return

    next_heading_index = _find_next_top_level_heading_index(paragraphs, title_index + 1)
    chart_paragraph = (
        paragraphs[title_index + 1]
        if title_index + 1 < next_heading_index
        else _insert_paragraph_after(paragraphs[title_index])
    )
    _clear_paragraph(chart_paragraph)
    chart_paragraph.add_run().add_picture(str(chart_path), width=Inches(5.8))

    for paragraph in paragraphs[title_index + 1 : next_heading_index]:
        if paragraph is not chart_paragraph:
            _remove_paragraph(paragraph)


def _replace_section_body(document: "DocxDocument", heading_text: str, replacement: str) -> None:
    paragraphs = document.paragraphs
    heading_index = _find_paragraph_index(paragraphs, heading_text)
    if heading_index is None:
        return

    next_heading_index = _find_next_top_level_heading_index(paragraphs, heading_index + 1)
    replacement_paragraph = (
        paragraphs[heading_index + 1]
        if heading_index + 1 < next_heading_index
        else _insert_paragraph_after(paragraphs[heading_index])
    )
    _set_paragraph_text(replacement_paragraph, replacement)

    for paragraph in paragraphs[heading_index + 1 : next_heading_index]:
        if paragraph is not replacement_paragraph:
            _remove_paragraph(paragraph)


def _set_paragraph_text(paragraph, text: str) -> None:
    _clear_paragraph(paragraph)
    if "\n" not in text:
        paragraph.add_run(text)
        return
    lines = text.splitlines()
    paragraph.add_run(lines[0])
    for line in lines[1:]:
        paragraph.add_run().add_break()
        paragraph.add_run(line)


def _bullet_text(items: list[str]) -> str:
    return "\n".join(f"{index}. {item}" for index, item in enumerate(items, start=1))


def _find_paragraph_index(paragraphs: list["Paragraph"], text: str) -> int | None:
    for index, paragraph in enumerate(paragraphs):
        if text in paragraph.text:
            return index
    return None


def _find_next_top_level_heading_index(paragraphs: list["Paragraph"], start_index: int) -> int:
    for index in range(start_index, len(paragraphs)):
        if re.match(r"^[一二三四五六七八九十]+、", paragraphs[index].text.strip()):
            return index
    return len(paragraphs)


def _insert_paragraph_after(paragraph: "Paragraph") -> "Paragraph":
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    return Paragraph(new_element, paragraph._parent)


def _clear_paragraph(paragraph: "Paragraph") -> None:
    _drop_unused_image_relationships(paragraph)
    paragraph.clear()


def _remove_paragraph(paragraph: "Paragraph") -> None:
    _drop_unused_image_relationships(paragraph)
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def _drop_unused_image_relationships(paragraph: "Paragraph") -> None:
    from docx.oxml.ns import qn

    embed_attr = qn("r:embed")
    relationship_ids = {
        blip.get(embed_attr)
        for blip in paragraph._element.xpath(".//a:blip")
        if blip.get(embed_attr)
    }
    for relationship_id in relationship_ids:
        references = [
            blip
            for blip in paragraph.part.element.xpath(".//a:blip")
            if blip.get(embed_attr) == relationship_id
        ]
        if len(references) == 1:
            paragraph.part.drop_rel(relationship_id)


def _report_period(metrics: MetricsResult) -> str:
    if metrics.paid_at_min and metrics.paid_at_max:
        return f"{metrics.paid_at_min:%Y-%m-%d} 至 {metrics.paid_at_max:%Y-%m-%d}"
    return "未识别"
